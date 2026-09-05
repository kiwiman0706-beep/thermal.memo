"""PDF / Word / テキスト / 画像ファイルの取り込み。

各ファイルから
  * テキスト抽出 (extract_text)
  * サムネイル画像 (render_pages)
のどちらかを選べるようにする。
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".log", ".json", ".yaml", ".yml"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"}
PDF_SUFFIXES = {".pdf"}
WORD_SUFFIXES = {".docx"}
OFFICE_SUFFIXES = {".doc", ".xlsx", ".xls", ".pptx", ".ppt", ".odt", ".rtf"}

SUPPORTED = TEXT_SUFFIXES | IMAGE_SUFFIXES | PDF_SUFFIXES | WORD_SUFFIXES | OFFICE_SUFFIXES


class DocumentError(RuntimeError):
    pass


@dataclass
class DocumentInfo:
    path: Path
    kind: str          # text / image / pdf / word / office
    pages: int
    can_text: bool
    can_thumbnail: bool
    note: str = ""


def _has(module: str) -> bool:
    import importlib.util

    return importlib.util.find_spec(module) is not None


def soffice_path() -> str | None:
    """LibreOffice / OpenOffice の CLI を探す（.doc/.xlsx のサムネイル用）。"""
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for candidate in (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if Path(candidate).exists():
            return candidate
    return None


def inspect(path: str | Path) -> DocumentInfo:
    path = Path(path)
    if not path.exists():
        raise DocumentError(f"ファイルが見つかりません: {path}")
    suffix = path.suffix.lower()

    if suffix in IMAGE_SUFFIXES:
        return DocumentInfo(path, "image", 1, can_text=_has("pytesseract"), can_thumbnail=True,
                            note="" if _has("pytesseract") else "テキスト抽出には pytesseract が必要")
    if suffix in PDF_SUFFIXES:
        pages = _pdf_page_count(path)
        thumb = _has("fitz")
        return DocumentInfo(path, "pdf", pages, can_text=_has("fitz") or _has("pdfplumber"),
                            can_thumbnail=thumb,
                            note="" if thumb else "サムネイルには PyMuPDF が必要")
    if suffix in WORD_SUFFIXES:
        return DocumentInfo(path, "word", 1, can_text=_has("docx"),
                            can_thumbnail=bool(soffice_path()),
                            note="" if _has("docx") else "テキスト抽出には python-docx が必要")
    if suffix in OFFICE_SUFFIXES:
        available = bool(soffice_path())
        return DocumentInfo(path, "office", 1, can_text=available, can_thumbnail=available,
                            note="" if available else "LibreOffice が見つかりません（PDF 化して読み込んでください）")
    if suffix in TEXT_SUFFIXES:
        return DocumentInfo(path, "text", 1, can_text=True, can_thumbnail=False)
    raise DocumentError(f"未対応の形式です: {suffix}")


def _pdf_page_count(path: Path) -> int:
    try:
        import fitz

        with fitz.open(path) as doc:
            return doc.page_count
    except ImportError:
        pass
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return len(pdf.pages)
    except Exception:
        return 1


# --------------------------------------------------------------------- テキスト

def extract_text(path: str | Path, pages: str = "all", ocr: bool = False) -> str:
    info = inspect(path)
    if info.kind == "text":
        return Path(path).read_text(encoding="utf-8", errors="replace")
    if info.kind == "pdf":
        text = _pdf_text(Path(path), pages)
        if not text.strip() and ocr:
            return _ocr_images(render_pages(path, pages=pages, dpi=200, width_dots=1200))
        return text
    if info.kind == "word":
        return _docx_text(Path(path))
    if info.kind == "office":
        return _soffice_text(Path(path))
    if info.kind == "image":
        return _ocr_images([Image.open(path)])
    raise DocumentError(f"テキスト抽出に未対応: {info.kind}")


def parse_page_spec(spec: str, total: int) -> list[int]:
    """'1,3-5' のような指定を 0 始まりのページ番号リストへ。"""
    spec = (spec or "all").strip().lower()
    if spec in ("", "all", "*"):
        return list(range(total))
    result: list[int] = []
    for part in spec.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            start, _, end = part.partition("-")
            try:
                lo = max(1, int(start or 1))
                hi = min(total, int(end or total))
            except ValueError:
                continue
            result.extend(range(lo - 1, hi))
        else:
            try:
                index = int(part) - 1
            except ValueError:
                continue
            if 0 <= index < total:
                result.append(index)
    return sorted(set(result)) or list(range(total))


def _pdf_text(path: Path, pages: str) -> str:
    try:
        import fitz

        with fitz.open(path) as doc:
            indexes = parse_page_spec(pages, doc.page_count)
            return "\n\n".join(doc[i].get_text("text").strip() for i in indexes).strip()
    except ImportError:
        pass
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            indexes = parse_page_spec(pages, len(pdf.pages))
            return "\n\n".join((pdf.pages[i].extract_text() or "").strip() for i in indexes).strip()
    except ImportError as exc:
        raise DocumentError("PDF のテキスト抽出には PyMuPDF か pdfplumber が必要です") from exc


def _docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise DocumentError(".docx の読み込みには python-docx が必要です") from exc
    document = docx.Document(str(path))
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    # 連続する空行を 1 行に畳む
    out: list[str] = []
    for line in chunks:
        if line.strip() or (out and out[-1].strip()):
            out.append(line.rstrip())
    return "\n".join(out).strip()


def _soffice_text(path: Path) -> str:
    pdf = _convert_with_soffice(path)
    return _pdf_text(pdf, "all")


def _ocr_images(images: list[Image.Image]) -> str:
    try:
        import pytesseract
    except ImportError as exc:
        raise DocumentError("OCR には pytesseract と Tesseract 本体が必要です") from exc
    return "\n\n".join(
        pytesseract.image_to_string(im, lang="jpn+eng").strip() for im in images
    ).strip()


# ------------------------------------------------------------------- サムネイル

def render_pages(
    path: str | Path,
    pages: str = "all",
    dpi: int = 150,
    width_dots: int = 576,
    max_pages: int = 20,
) -> list[Image.Image]:
    """ページ画像のリストを返す（印刷幅に合わせて縮小済み・グレースケール）。"""
    path = Path(path)
    info = inspect(path)

    if info.kind == "image":
        images = [Image.open(path)]
    elif info.kind == "pdf":
        images = _pdf_pages(path, pages, dpi, max_pages)
    elif info.kind in ("word", "office"):
        images = _pdf_pages(_convert_with_soffice(path), pages, dpi, max_pages)
    else:
        raise DocumentError(f"サムネイル生成に未対応: {info.kind}")

    out: list[Image.Image] = []
    for image in images:
        gray = image.convert("L")
        if gray.width != width_dots:
            height = max(1, round(gray.height * width_dots / gray.width))
            gray = gray.resize((width_dots, height), Image.LANCZOS)
        out.append(gray)
    return out


def _pdf_pages(path: Path, pages: str, dpi: int, max_pages: int) -> list[Image.Image]:
    try:
        import fitz
    except ImportError as exc:
        raise DocumentError("PDF のサムネイルには PyMuPDF が必要です") from exc
    result: list[Image.Image] = []
    with fitz.open(path) as doc:
        indexes = parse_page_spec(pages, doc.page_count)[:max_pages]
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        for index in indexes:
            pix = doc[index].get_pixmap(matrix=matrix, colorspace=fitz.csGRAY)
            result.append(Image.frombytes("L", (pix.width, pix.height), pix.samples))
    return result


def _convert_with_soffice(path: Path) -> Path:
    soffice = soffice_path()
    if not soffice:
        raise DocumentError(
            "この形式のサムネイル生成には LibreOffice が必要です。"
            "PDF に書き出してから読み込んでください。"
        )
    outdir = Path(tempfile.mkdtemp(prefix="thermal_memo_"))
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(outdir), str(path)],
        check=False, timeout=120,
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )
    converted = outdir / (path.stem + ".pdf")
    if not converted.exists():
        raise DocumentError(f"LibreOffice による PDF 変換に失敗しました: {path.name}")
    return converted
